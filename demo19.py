from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading("大家好，老子叫马云")

document.add_paragraph("老子的淘宝什么东西都有卖，你们这些老表快点去买")

#document.add_page_break()

document.add_picture("F:\py\.idea\imgs\jackma.jpg",width=Inches(1.25))

document.save("test7.docx")